from __future__ import annotations

import datetime as _dt
from pathlib import Path

from docx import Document


def build_docx(out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)

    today = _dt.date.today().isoformat()

    doc = Document()

    doc.add_heading("SmartHub Diagnostics — Conceptual Mapping", level=1)
    doc.add_paragraph(f"Generated: {today}")
    doc.add_paragraph("Scope: USB-only BSOD triage + offline helpers")

    doc.add_heading("Concept Map (Text View)", level=2)
    doc.add_paragraph(
        "This is a DOCX-friendly text representation of the conceptual mapping (the PDF contains the visual bubble map)."
    )

    branches: dict[str, list[str]] = {
        "UI Layer": [
            "SmartHub.exe (WPF/WebView2)",
            "Web UI (HTML/JS/CSS)",
        ],
        "Backend API": [
            "Express (:3333)",
            "/connection-check",
            "Loopback 127.0.0.1",
        ],
        "Android Phone": [
            "USB cable",
            "MTP / ADB / Fastboot",
            "Screen: normal/dark",
        ],
        "Windows Evidence": [
            "PnP/CIM sampling",
            "Shell MTP probe",
            "WPD helper (optional)",
        ],
        "Offline Helpers": [
            "Camera/OCR check",
            "Offline AI (local-only)",
            "Supporting only",
        ],
        "Storage": [
            "history.json",
            "memory.sqlite",
            "logs/screenshots",
        ],
        "Truthfulness Rules": [
            "No USB enum → INCONCLUSIVE",
            "COM/WPD error → host-only",
            "Probe unavailable → ignore camera",
            "Manual “UI frozen” only",
        ],
    }

    for branch, items in branches.items():
        p = doc.add_paragraph(branch, style="List Bullet")
        p.runs[0].bold = True
        for item in items:
            doc.add_paragraph(item, style="List Bullet 2")

    doc.add_heading("Narrative", level=2)
    narrative_lines = [
        "1) Technician opens SmartHub (WPF/WebView2). The embedded Web UI runs locally.",
        "2) Web UI calls the local Node backend (127.0.0.1:3333), mainly /connection-check for USB-only triage.",
        "3) Backend collects Windows evidence (PnP/CIM USB sampling, optional WPD helper, Shell-based MTP probe fallback).",
        "4) Optional camera/OCR and offline AI are best-effort supporting tools; they must not override truthfulness rules.",
        "5) Safety rule: if the PC cannot enumerate the phone over USB, the result is INCONCLUSIVE.",
        "6) Host limitation rule: if WPD/COM is broken (E_NOINTERFACE), the app outputs host-inconclusive guidance and does not infer the phone’s state.",
        "   Camera hints are ignored; only manual “UI frozen” confirmation can record freeze.",
    ]
    for line in narrative_lines:
        doc.add_paragraph(line)

    doc.add_heading("Shape Meanings (for the PDF diagram)", level=2)
    shape_lines = [
        "Cloud (center) = the main concept/system being mapped.",
        "Large ovals = major components/actors (primary branches).",
        "Small ovals = subcomponents, signals, or rules (supporting details).",
        "Lines = conceptual relationships / information flow (not strict sequencing).",
    ]
    for line in shape_lines:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Connection Relations (Text Mode)", level=2)
    doc.add_paragraph("A → B means A calls/uses/sends data to B.")

    sections: list[tuple[str, list[str]]] = [
        (
            "UI ↔ Backend",
            [
                "SmartHub.exe (WPF/WebView2) → Web UI (HTML/JS/CSS): hosts/embeds the UI.",
                "Web UI → Node Backend (Express :3333): loopback HTTP requests (127.0.0.1).",
                "Node Backend → Web UI: JSON response used to render diagnostic status/result.",
            ],
        ),
        (
            "USB Evidence",
            [
                "Node Backend → Windows Host OS: PowerShell/CIM/PnP enumeration + USB stability sampling.",
                "Node Backend → Shell MTP probe: best-effort check for MTP accessibility without WPD COM dependency.",
                "Node Backend → UsbEvidenceHelper (WPD/COM, optional): deeper MTP/WPD probing when available.",
                "Windows Host OS ↔ Android Phone: USB cable enumeration (MTP/ADB/Fastboot interfaces when present).",
            ],
        ),
        (
            "Offline Helpers",
            [
                "Node Backend → Camera/OCR helper: optional visual hints (screen dialog / darkness) when supported.",
                "Node Backend → Offline AI helper: optional local-only suggestion based on collected signals.",
            ],
        ),
        (
            "Storage",
            [
                "Node Backend → Local artifacts: writes logs, run history, and (optional) screenshots/AI cases.",
                "Web UI → Local storage/history view: reads prior runs for technician review (implementation-specific).",
            ],
        ),
        (
            "Truthfulness Gates (Always Apply)",
            [
                "No Windows USB enumeration → output INCONCLUSIVE (do not claim phone BSOD/freeze).",
                "Host COM/WPD error → host-inconclusive (do not infer phone state; ignore camera hints).",
                "Manual “UI frozen” checkbox → only way to record freeze when probing is unavailable.",
            ],
        ),
    ]

    for title, items in sections:
        doc.add_paragraph(title, style="List Bullet")
        for item in items:
            doc.add_paragraph(item, style="List Bullet 2")

    doc.add_paragraph("Regenerate this DOCX after major flow changes.")

    doc.save(str(out_path))


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    out = repo_root / "docx" / "SmartHub_Conceptual_Mapping.docx"
    build_docx(out)
    print(str(out))


if __name__ == "__main__":
    main()
