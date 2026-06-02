from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_arrows_table(doc: Document, arrows: list[tuple[str, str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    header = table.rows[0].cells
    header[0].text = "From"
    header[1].text = "To"
    header[2].text = "Arrow Text"
    header[3].text = "What It Does"

    for src, dst, label, meaning in arrows:
        row = table.add_row().cells
        row[0].text = src
        row[1].text = dst
        row[2].text = label
        row[3].text = meaning


def add_level_section(
    doc: Document,
    level_title: str,
    users: list[str],
    processes: list[str],
    databases: list[str],
    arrows: list[tuple[str, str, str, str]],
) -> None:
    doc.add_heading(level_title, level=2)

    doc.add_heading("Users / External Entities", level=3)
    add_bullets(doc, users)

    doc.add_heading("Processes", level=3)
    add_bullets(doc, processes)

    doc.add_heading("Database / Data Stores", level=3)
    add_bullets(doc, databases)

    doc.add_heading("Data Flow Arrows", level=3)
    add_arrows_table(doc, arrows)


def build_doc(out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("SmartHub Diagnostics - Gane-Sarson DFD (Level 0-2)", level=1)
    doc.add_paragraph(f"Generated: {date.today().isoformat()}")
    doc.add_paragraph(
        "This document lists users, processes, databases/data stores, and labeled arrows "
        "for DFD Level 0, Level 1, and Level 2 using Gane-Sarson style notation."
    )

    add_level_section(
        doc,
        "DFD Level 0 (Context Diagram)",
        users=[
            "U1 Technician",
            "U2 Android Phone",
            "U3 Supabase Auth Service",
        ],
        processes=[
            "P0 SmartHub Diagnostic App",
        ],
        databases=[
            "D1 Local Storage: history.json, auth-local-session.json, memory.sqlite, adb_ai_memory.sqlite",
            "D2 Cloud Storage (optional): Supabase diagnostic_runs",
        ],
        arrows=[
            ("Technician", "SmartHub Diagnostic App", "Login", "Technician enters credentials to start an authenticated session."),
            ("SmartHub Diagnostic App", "Supabase Auth Service", "Verify Credentials", "App submits login credentials for validation."),
            ("Supabase Auth Service", "SmartHub Diagnostic App", "Auth Result", "Returns success/failure plus user identity."),
            ("SmartHub Diagnostic App", "Technician", "Login Status", "App shows sign-in result and access state."),
            ("SmartHub Diagnostic App", "Android Phone", "Diagnostic Command", "App sends ADB/USB diagnostic commands."),
            ("Android Phone", "SmartHub Diagnostic App", "Device Evidence", "Phone returns logs, status, and test outputs."),
            ("SmartHub Diagnostic App", "D1 Local Storage", "Save Local Run", "Stores run result, local AI memory, and local auth session."),
            ("SmartHub Diagnostic App", "D2 Supabase diagnostic_runs", "Cloud Sync", "Mirrors diagnostic run for the authenticated owner user."),
            ("D2 Supabase diagnostic_runs", "SmartHub Diagnostic App", "Cloud History", "Loads prior user-owned diagnostic runs."),
        ],
    )

    add_level_section(
        doc,
        "DFD Level 1 (Major Process Decomposition)",
        users=[
            "U1 Technician",
            "U2 Android Phone",
            "U3 Supabase Auth Service",
        ],
        processes=[
            "P1 Authenticate Technician",
            "P2 Discover Device Connection",
            "P3 Run Diagnostics",
            "P4 Generate AI Conclusion",
            "P5 Save and Sync Results",
        ],
        databases=[
            "D1 Local Storage: history.json, auth-local-session.json, memory.sqlite, adb_ai_memory.sqlite",
            "D2 Supabase diagnostic_runs",
        ],
        arrows=[
            ("Technician", "P1 Authenticate Technician", "Login", "Technician submits email/username and password."),
            ("P1 Authenticate Technician", "Supabase Auth Service", "Validate Login", "Checks credentials and obtains identity."),
            ("P1 Authenticate Technician", "D1 Local Storage", "Write Offline Session", "Stores auth-local-session token for offline reuse."),
            ("Technician", "P2 Discover Device Connection", "Start Check", "Technician starts connection check and selects target device."),
            ("P2 Discover Device Connection", "Android Phone", "Probe Device", "Collects ADB/USB/PnP connection evidence."),
            ("P2 Discover Device Connection", "P3 Run Diagnostics", "Connection Evidence", "Passes validated connection signals to diagnostics."),
            ("P3 Run Diagnostics", "P4 Generate AI Conclusion", "Feature Summary", "Sends extracted findings for AI suggestion."),
            ("P4 Generate AI Conclusion", "P5 Save and Sync Results", "AI Recommendation", "Returns likely cause and next actions."),
            ("P5 Save and Sync Results", "D1 Local Storage", "Save Run", "Writes local run history and artifacts."),
            ("P5 Save and Sync Results", "D2 Supabase diagnostic_runs", "Save Cloud Run", "Writes owner-scoped cloud diagnostic record."),
            ("P5 Save and Sync Results", "Technician", "Show Result", "Displays final diagnosis and recommendations."),
        ],
    )

    add_level_section(
        doc,
        "DFD Level 2 (Detailed Decomposition of P1 Authenticate Technician)",
        users=[
            "U1 Technician",
            "U3 Supabase Auth Service",
        ],
        processes=[
            "P1.1 Capture Login Input",
            "P1.2 Validate Credentials Online",
            "P1.3 Create Offline Session",
            "P1.4 Return Auth Response",
        ],
        databases=[
            "D1 auth-local-session.json",
            "D3 Supabase auth.users (managed service)",
        ],
        arrows=[
            ("Technician", "P1.1 Capture Login Input", "Login", "Technician provides login credentials in SmartHub UI."),
            ("P1.1 Capture Login Input", "P1.2 Validate Credentials Online", "Credentials", "Passes sanitized credentials for online verification."),
            ("P1.2 Validate Credentials Online", "Supabase Auth Service", "Auth Request", "Requests sign-in validation against Supabase auth."),
            ("Supabase Auth Service", "P1.2 Validate Credentials Online", "Auth Response", "Returns user id, email, and authentication status."),
            ("P1.2 Validate Credentials Online", "P1.3 Create Offline Session", "Authenticated User", "Provides verified identity to create local offline token."),
            ("P1.3 Create Offline Session", "D1 auth-local-session.json", "Store Session", "Writes token, user id, email, and expiry for offline use."),
            ("P1.3 Create Offline Session", "P1.4 Return Auth Response", "Session Metadata", "Passes token and expiration metadata to response handler."),
            ("P1.4 Return Auth Response", "Technician", "Login Success/Failure", "Returns auth result and next-step message in UI."),
        ],
    )

    doc.add_heading("Drawing Note", level=2)
    doc.add_paragraph(
        "When drawing in Gane-Sarson format: use rectangles for users/external entities, "
        "process symbols for P-nodes, open-ended stores for D-nodes, and labeled arrows "
        "using the exact arrow text listed above."
    )

    try:
        doc.save(out_path)
        print(f"Wrote: {out_path}")
    except PermissionError:
        fallback = out_path.with_name("Gane-Sarson-DFD-Level0-2.updated.docx")
        doc.save(fallback)
        print(f"Target file is locked. Wrote: {fallback}")


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    out = repo_root / "docx" / "Gane-Sarson-DFD-Level0-2.docx"
    build_doc(out)


if __name__ == "__main__":
    main()
