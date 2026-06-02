from __future__ import annotations

from pathlib import Path

from docx import Document


def _clean(text: str) -> str:
    return " ".join(text.replace("\n", " ").split()).strip()


def main() -> None:
    doc_path = Path("docx") / "Chapter 3-smarthub.docx"
    doc = Document(doc_path)

    print(f"Loaded: {doc_path}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")

    for ti, table in enumerate(doc.tables):
        headers: list[str] = []
        if table.rows:
            headers = [_clean(c.text) for c in table.rows[0].cells]
        print("\n" + ("=" * 60))
        print(f"Table {ti}: rows={len(table.rows)} cols={len(table.columns)}")
        print(f"Headers: {headers}")

        # show a few sample rows
        for ri in range(1, min(8, len(table.rows))):
            row = table.rows[ri]
            texts = [_clean(c.text) for c in row.cells]
            print(f"  r{ri}: {texts}")

        # quick scan for description-like column and blanks
        desc_cols = [i for i, h in enumerate(headers) if h.lower() in {"description", "meaning", "details"}]
        if desc_cols:
            dc = desc_cols[0]
            blanks = 0
            for ri in range(1, len(table.rows)):
                cell_text = _clean(table.rows[ri].cells[dc].text)
                if cell_text in {"", "tbd", "n/a", "na", "-"}:
                    blanks += 1
            print(f"Possible description column index={dc}; blanks={blanks}")


if __name__ == "__main__":
    main()
