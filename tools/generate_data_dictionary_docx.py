from __future__ import annotations

from pathlib import Path
from typing import TypedDict

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_FILES = [
    ROOT / "docx" / "data-dictionary.docx",
    ROOT / "docx" / "data-dictionary.updated.docx",
]

HEADERS = ["Data Name", "Field Size", "Data Type", "Data Format", "Description", "Example"]
COLUMN_WIDTHS = [Inches(1.5), Inches(1.0), Inches(1.1), Inches(1.2), Inches(1.8), Inches(1.1)]


class SectionDef(TypedDict):
    title: str
    definition: str
    rows: list[tuple[str, str, str, str, str, str]]


def shade_cell(cell, fill: str = "D9D9D9") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, center: bool = False) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if center else WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        if not paragraph.runs:
            paragraph.add_run("")
        for run in paragraph.runs:
            run.bold = bold
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)


def set_paragraph_style(
    paragraph,
    *,
    align: WD_PARAGRAPH_ALIGNMENT = WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
    bold: bool = False,
    underline: bool = False,
    size: int = 11,
    space_after: int = 6,
) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.space_before = Pt(0)
    for run in paragraph.runs:
        run.bold = bold
        run.underline = underline
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)


def comma_join(values: list[str]) -> str:
    if not values:
        return ""
    if len(values) == 1:
        return values[0]
    if len(values) == 2:
        return f"{values[0]} and {values[1]}"
    return f"{', '.join(values[:-1])}, and {values[-1]}"


def list_join(values: list[str]) -> str:
    return ", ".join(values)


def number_word(n: int) -> str:
    words = {
        0: "zero",
        1: "one",
        2: "two",
        3: "three",
        4: "four",
        5: "five",
        6: "six",
        7: "seven",
        8: "eight",
        9: "nine",
        10: "ten",
        11: "eleven",
        12: "twelve",
    }
    return words.get(n, str(n))


def ordered_type_signature(section: SectionDef) -> str:
    seen: list[str] = []
    for row in section["rows"]:
        dtype = str(row[2]).strip().lower()
        if dtype and dtype not in seen:
            seen.append(dtype)
    return "/".join(seen)


def criticality_clause(section: SectionDef) -> str:
    title = section["title"]
    clauses: dict[str, str] = {
        "password_reset_token": (
            "it secures account recovery flows by validating reset tokens and timestamp windows, reducing unauthorized password resets"
        ),
        "sessions": (
            "it tracks authenticated activity, device access context, and session continuity for secure and traceable user access"
        ),
        "user": (
            "it centralizes account identity and credential metadata required for authentication, confirmation, and lifecycle management"
        ),
        "app_user": (
            "it links identity records with profile attributes so user level ownership, display context, and communication details remain consistent"
        ),
        "device": (
            "it preserves stable device identity and visibility windows, enabling accurate device mapping, auditability, and longitudinal diagnostics"
        ),
        "diagnostic_runs": (
            "it links users, devices, and diagnostic types into a single session record, enabling accurate history tracking, performance analysis, and cloud based reporting across all diagnostic runs"
        ),
        "result": (
            "it binds computed outcomes to diagnostic and AI references, supporting dependable interpretation trails and reproducible technical reporting"
        ),
        "smarthub_ai": (
            "it stores model generated conclusions and payload evidence used to explain, compare, and improve SmartHub diagnostic intelligence"
        ),
    }
    return clauses.get(
        title,
        "it captures structured diagnostic context required for reliable tracking, analysis, and reporting",
    )


def table_intro_text(number: int, section: SectionDef) -> str:
    field_names = [row[0] for row in section["rows"]]
    field_count = len(field_names)
    field_types = ordered_type_signature(section)
    table_name = section["title"]
    table_label = table_name.replace("_", " ")

    return (
        f"Table 3.{number} is entry for the {table_name} table documents "
        f"{number_word(field_count)} fields ({list_join(field_names)}) with their sizes, "
        f"{field_types} types, formats, descriptions, and examples; according to ISO/IEC 25010 "
        "(software quality model), this specification ensures functional suitability by defining "
        f"the {table_label} structure, reliability via timestamp and foreign key constraints, and "
        "maintainability through clear field and payload definitions; moreover, this entry is "
        f"critical for SmartHub because {criticality_clause(section)}."
    )


def add_section(document: Document, number: int, section: SectionDef) -> None:
    intro = document.add_paragraph(table_intro_text(number, section))
    set_paragraph_style(intro, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, size=11, space_after=6)

    define = document.add_paragraph(f"Definition: {section['definition']}")
    set_paragraph_style(define, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, size=11, space_after=6)

    label = document.add_paragraph(f"Table 3.{number} {section['title']}")
    set_paragraph_style(label, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, bold=True, size=11, space_after=4)

    table = document.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"
    table.autofit = False

    header_cells = table.rows[0].cells
    for idx, header in enumerate(HEADERS):
        header_cells[idx].width = COLUMN_WIDTHS[idx]
        set_cell_text(header_cells[idx], header, bold=True, center=True)
        shade_cell(header_cells[idx])

    for row in section["rows"]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = COLUMN_WIDTHS[idx]
            set_cell_text(cells[idx], value)

    document.add_paragraph("")


def erd_sections() -> list[SectionDef]:
    return [
        {
            "title": "password_reset_token",
            "definition": "Stores user password reset token records and creation timestamps for account recovery.",
            "rows": [
                ("email", "255", "text", "email", "Foreign key email reference to user", "user@example.com"),
                ("token", "255", "text", "token", "Password reset token value", "4xZs9u..."),
                ("created_at", "8 bytes", "timestamptz", "datetime", "Token creation timestamp", "2026-04-19 10:10:12+00"),
            ],
        },
        {
            "title": "sessions",
            "definition": "Stores authenticated session records linked to each user account.",
            "rows": [
                ("id", "255", "text", "session id", "Primary key of session", "s1cf4d8a..."),
                ("user_id", "36", "uuid", "UUID", "Foreign key to user.user_id", "f04d0dc8-7fef-4f8b-9d8b-1f0a12345678"),
                ("ip_address", "45", "text", "IPv4/IPv6", "Client IP address", "192.168.1.15"),
                ("user_agent", "variable", "text", "browser/agent", "Client user-agent string", "SmartHubDesktop/14.0"),
                ("payload", "variable", "text", "serialized", "Session payload/body", "{...}"),
                ("last_activity", "8 bytes", "bigint", "epoch", "Last activity epoch time", "1713525000"),
            ],
        },
        {
            "title": "user",
            "definition": "Stores authenticated user account credentials and registration metadata.",
            "rows": [
                ("user_id", "36", "uuid", "UUID", "Primary key user identifier", "f04d0dc8-7fef-4f8b-9d8b-1f0a12345678"),
                ("email", "255", "text", "email", "User email", "user@example.com"),
                ("encrypted_password", "255", "text", "hash", "Stored password hash", "$2a$12$..."),
                ("email_confirmed_at", "8 bytes", "timestamptz", "datetime", "Email confirmation time", "2026-04-19 09:55:00+00"),
                ("created_at", "8 bytes", "timestamptz", "datetime", "User creation timestamp", "2026-04-19 09:50:00+00"),
            ],
        },
        {
            "title": "app_user",
            "definition": "Stores profile metadata linked to each authenticated user account.",
            "rows": [
                ("owner_user_id", "36", "uuid", "UUID", "Primary key for app user", "f04d0dc8-7fef-4f8b-9d8b-1f0a12345678"),
                ("user_id", "36", "uuid", "UUID", "Foreign key to user.user_id", "f04d0dc8-7fef-4f8b-9d8b-1f0a12345678"),
                ("display_name", "120", "text", "text", "Display name", "Mark Lhemuel"),
                ("email", "255", "text", "email", "User email copy", "user@example.com"),
                ("created_at", "8 bytes", "timestamptz", "datetime", "Profile creation timestamp", "2026-04-19 09:51:00+00"),
            ],
        },
        {
            "title": "device",
            "definition": "Stores tracked Android device identifiers and first/last seen timestamps.",
            "rows": [
                ("device_id", "64", "text", "device id", "Primary key device identifier", "R58M12345AB"),
                ("device_label", "120", "text", "text", "Human-readable device label", "Samsung A14 Bench Unit"),
                ("first_seen_at", "8 bytes", "timestamptz", "datetime", "First seen timestamp", "2026-04-19 09:40:00+00"),
                ("last_seen_at", "8 bytes", "timestamptz", "datetime", "Most recent seen timestamp", "2026-04-19 10:30:00+00"),
            ],
        },
        {
            "title": "diagnostic_runs",
            "definition": "Stores each diagnostic execution record linked to the app user and device.",
            "rows": [
                ("diagnostic_id", "8 bytes", "bigint", "integer", "Primary key diagnostic identifier", "1713525123"),
                ("owner_user_id", "36", "uuid", "UUID", "Foreign key to app_user.owner_user_id", "f04d0dc8-7fef-4f8b-9d8b-1f0a12345678"),
                ("device_id", "64", "text", "device id", "Foreign key to device.device_id", "R58M12345AB"),
                ("diagnostic_type", "64", "text", "text", "Type of diagnostic", "usb-no-debug"),
                ("run_timestamp", "8 bytes", "bigint", "epoch ms", "Diagnostic timestamp", "1713525123456"),
                ("payload", "variable", "jsonb", "JSON object", "Diagnostic payload", "{\"diagStages\":{...}}"),
            ],
        },
        {
            "title": "result",
            "definition": "Stores output results linked to diagnostic runs and AI-generated conclusions.",
            "rows": [
                ("result_id", "8 bytes", "bigint", "integer", "Primary key result identifier", "880001"),
                ("diagnostic_id", "8 bytes", "bigint", "integer", "Foreign key to diagnostic_runs.diagnostic_id", "1713525123"),
                ("ai_id", "8 bytes", "bigint", "integer", "Foreign key to smarthub_ai.ai_id", "990010"),
                ("result", "variable", "text", "text", "Final diagnostic result text", "Likely display controller fault"),
                ("run_timestamp", "8 bytes", "bigint", "epoch ms", "Result creation timestamp", "1713525123900"),
                ("payload", "variable", "jsonb", "JSON object", "Detailed result payload", "{\"findings\":[...] }"),
            ],
        },
        {
            "title": "smarthub_ai",
            "definition": "Stores AI inference outputs used by SmartHub to support diagnostic conclusions.",
            "rows": [
                ("ai_id", "8 bytes", "bigint", "integer", "Primary key AI record identifier", "990010"),
                ("conclusion", "variable", "text", "text", "AI-generated conclusion summary", "High probability of software-level UI freeze"),
                ("run_timestamp", "8 bytes", "bigint", "epoch ms", "Inference timestamp", "1713525123888"),
                ("payload", "variable", "jsonb", "JSON object", "Full AI response payload", "{\"model\":\"gpt-4o-mini\",...}"),
            ],
        },
    ]


def generate_document() -> Document:
    doc = Document()

    page = doc.sections[0]
    page.top_margin = Inches(0.7)
    page.bottom_margin = Inches(0.7)
    page.left_margin = Inches(0.6)
    page.right_margin = Inches(0.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    figure_title = doc.add_paragraph("Figure 3.5.9 - Data Dictionary")
    set_paragraph_style(
        figure_title,
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        bold=True,
        underline=True,
        size=13,
        space_after=8,
    )

    intro = doc.add_paragraph(
        "This section states and defines the key Supabase ERD tables and fields used by SmartHub "
        "to store user, session, device, diagnostic, AI, and result records. "
        "All narrative entries are justified and table entries are numbered for documentation consistency."
    )
    set_paragraph_style(intro, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, size=11, space_after=10)

    sections = erd_sections()
    for index, section in enumerate(sections, start=1):
        add_section(doc, index, section)

    return doc


def main() -> None:
    document = generate_document()
    for output in OUTPUT_FILES:
        output.parent.mkdir(parents=True, exist_ok=True)
        document.save(output)
        print(f"Updated {output}")


if __name__ == "__main__":
    main()